"""imprint/srs.py — render a requirement set as a Software Requirements
Specification (.docx), following ISO/IEC/IEEE 29148 structure.

"Requirements are data, not a document" — this module is the first *view* of that
data: it presses the stored requirement records into a formatted SRS. Built on
python-docx (already a Book Reader dependency), so it works 100% offline with no
template file to manage. UI-free, so it's unit-tested without a window.
"""

from __future__ import annotations

import os
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from . import models


def default_output_path(project_name: str, root: str | None = None) -> str:
    """A sensible save location: <root>/output/<Project>_SRS.docx."""
    root = root or os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    safe = re.sub(r"[^\w\- ]+", "", project_name).strip().replace(" ", "_") or "Project"
    return os.path.join(root, "output", f"{safe}_SRS.docx")


def _labeled(doc, label: str, value: str) -> None:
    """Add a 'Label: value' paragraph with the label in bold (skips empty values)."""
    value = (value or "").strip()
    if not value:
        return
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)


def build_srs(project, requirements, out_path: str, generated_on: str | None = None) -> str:
    """Write an SRS .docx for `project` + its `requirements`; return the path.

    `project` and `requirements` are sqlite3.Row (dict-like) records from db.py.
    """
    generated_on = generated_on or datetime.now().strftime("%Y-%m-%d")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    doc = Document()
    doc.core_properties.title = f"{project['name']} — Software Requirements Specification"

    # --- title block ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(project["name"])
    r.bold = True
    r.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Software Requirements Specification")
    sr.font.size = Pt(14)

    std = doc.add_paragraph()
    std.alignment = WD_ALIGN_PARAGRAPH.CENTER
    std.add_run("Prepared with Imprint · targeting ISO/IEC/IEEE 29148").italic = True

    # --- document information table ---
    info = [
        ("Project", project["name"]),
        ("Methodology", models.methodology_label(project["methodology"])),
        ("Generated", generated_on),
        ("Requirements", str(len(requirements))),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(info):
        c0, c1 = table.rows[i].cells
        c0.paragraphs[0].add_run(k).bold = True
        c1.text = v

    # --- 1. Introduction ---
    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        f"This document specifies the software requirements for {project['name']}. "
        "It records each requirement as a uniquely identified, verifiable statement "
        "so the system can be built, tested, and traced against it."
    )
    doc.add_heading("1.2 Scope", level=2)
    doc.add_paragraph((project["description"] or "").strip()
                      or "[Scope to be completed — describe what the system will and will not do.]")

    # --- 2. Specific Requirements (grouped by type) ---
    doc.add_heading("2. Specific Requirements", level=1)
    if not requirements:
        doc.add_paragraph("[No requirements captured yet.]")
    else:
        section = 1
        for rtype in models.REQ_TYPES:
            group = [q for q in requirements if q["req_type"] == rtype]
            if not group:
                continue
            doc.add_heading(f"2.{section} {rtype} Requirements", level=2)
            section += 1
            for q in group:
                doc.add_heading(f"{q['req_key']} — {q['moscow']}", level=3)
                doc.add_paragraph(q["statement"])
                _labeled(doc, "Acceptance criteria", q["acceptance_criteria"])
                _labeled(doc, "Source", q["source"])
                _labeled(doc, "Rationale", q["rationale"])
                _labeled(doc, "Status", q["status"])

    # --- Appendix: requirements summary ---
    if requirements:
        doc.add_heading("Appendix A — Requirements Summary", level=1)
        summary = doc.add_table(rows=1, cols=4)
        summary.style = "Table Grid"
        for cell, head in zip(summary.rows[0].cells, ("ID", "Type", "Priority", "Statement")):
            cell.paragraphs[0].add_run(head).bold = True
        for q in requirements:
            cells = summary.add_row().cells
            cells[0].text = q["req_key"]
            cells[1].text = q["req_type"]
            cells[2].text = q["moscow"]
            cells[3].text = q["statement"]

    doc.save(out_path)
    return out_path
