"""imprint/user_stories.py — render requirements as Agile user stories (.docx).

The third view of the same requirement data (Waterfall = SRS, V-Model = matrix,
Agile = this). A stored "The system shall ..." requirement is rewritten into the
"As a user, I want ... so that ..." story form and grouped by MoSCoW priority
(the Agile backlog ordering). Deterministic — no model needed, works offline.
UI-free, so it's unit-tested without a window.
"""

from __future__ import annotations

import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from . import models

_PREFIX = "the system shall "


def to_user_story(statement: str, rationale: str = "") -> str:
    """Rewrite a 'The system shall ...' requirement as an Agile user story (pure)."""
    cap = (statement or "").strip().rstrip(".")
    if not cap:
        return "As a user, I want ..."
    if cap.lower().startswith(_PREFIX):
        cap = cap[len(_PREFIX):].strip()
        story = f"As a user, I want to {cap}"
    else:
        story = f"As a user, I want {cap[0].lower() + cap[1:]}"
    if (rationale or "").strip():
        story += f", so that {rationale.strip().rstrip('.')}"
    return story + "."


def default_output_path(project_name: str, root: str | None = None) -> str:
    root = root or os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    safe = re.sub(r"[^\w\- ]+", "", project_name).strip().replace(" ", "_") or "Project"
    return os.path.join(root, "output", f"{safe}_User_Stories.docx")


def build_user_stories(project, requirements, out_path: str) -> str:
    """Write a user-stories .docx for `project` + its `requirements`; return the path."""
    os.makedirs(os.path.dirname(out_path), exist_ok=True)

    doc = Document()
    doc.core_properties.title = f"{project['name']} — User Stories"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(project["name"])
    r.bold = True
    r.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("User Stories (Agile)").font.size = Pt(14)

    doc.add_paragraph(
        "The same requirements, expressed as prioritized user stories for backlog "
        "and sprint planning. Grouped by MoSCoW priority."
    )

    if not requirements:
        doc.add_paragraph("[No requirements captured yet.]")
    else:
        for priority in models.MOSCOW:  # Must -> Should -> Could -> Won't
            group = [q for q in requirements if q["moscow"] == priority]
            if not group:
                continue
            doc.add_heading(f"{priority} have ({len(group)})", level=1)
            for q in group:
                doc.add_heading(q["req_key"], level=3)
                doc.add_paragraph(to_user_story(q["statement"], q["rationale"]))
                if (q["acceptance_criteria"] or "").strip():
                    label = doc.add_paragraph()
                    label.add_run("Acceptance criteria: ").bold = True
                    label.add_run(q["acceptance_criteria"].strip())
                status = doc.add_paragraph()
                status.add_run("Status: ").bold = True
                status.add_run(q["status"])

    doc.save(out_path)
    return out_path
